#  Author: Anton Nedilko
#   Email: arcs3567@gmail.com
#  Source: https://github.com/techeditorsrc
#    Info: resume builder
# License: GPL v3

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.dml.color import ColorFormat
from docx.shared import Pt
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
import os

current_dir=os.getcwd().replace('\\','/')+'/'

def int_to_hex(p,align=-1):
    r=hex(p)[2:]
    if align!=-1:
        l=len(r)
        if l<align:
            r='0'*(align-l)+r
    return r

def rgb_to_hex(r,g,b):
    return int_to_hex(r,align=2)+int_to_hex(g,align=2)+int_to_hex(b,align=2)

def copy(p):
    r=p
    if isinstance(p,tuple):
        r=[]
        for i in p:
            r.append(copy(i))
        r=tuple(r)
        return r
    elif isinstance(p,list):
        r=[]
        for i in p:
            r.append(copy(i))
        return r
    elif isinstance(p,dict):
        r={}
        for key,value in p.items():
            r[key]=copy(value)
        return r
    else:
        return r
class resume():
    def new_font_style(self,font_name,font_size,font_bold,font_italic,font_underline,font_color):
        return {'name':font_name,'size':font_size,'bold':font_bold,'italic':font_italic,'underline':font_underline,'color':font_color}
    
    def new_paragraph_style(self,space_before,space_after):
        return {'space_before':space_before,'space_after':space_after}

    def add_param(self,style_name,param_name,params):
        if not style_name in self.styles:
            self.styles[style_name]={}
        if not param_name in self.styles[style_name]:
            self.styles[style_name][param_name]={}
        self.styles[style_name][param_name]=params

    def default_style(self):
        style_name='default'
        font_size=10
        font_name='Dejavu Sans'
        default_color=(0,0,0)
        section_color=(65,138,183)
        paragraph_title_color=(53,108,132)
        name_style={'font':self.new_font_style(font_name,font_size,True,False,False,default_color),'paragraph':self.new_paragraph_style(2,2)}
        speciality_style={'font':self.new_font_style(font_name,font_size,False,False,False,default_color),'paragraph':self.new_paragraph_style(2,2)}
        contacts_style={'font':self.new_font_style(font_name,font_size-2,False,False,False,default_color),'paragraph':self.new_paragraph_style(2,2)}
        link_style={'font':self.new_font_style(font_name,font_size-2,False,False,False,section_color),'paragraph':self.new_paragraph_style(2,2)}
        section_name_style={'font':self.new_font_style(font_name,font_size,False,False,False,section_color),'paragraph':self.new_paragraph_style(2,2)}
        paragraph_name_style={'font':self.new_font_style(font_name,font_size,False,False,False,paragraph_title_color),'paragraph':self.new_paragraph_style(2,2)}
        paragraph_text_style={'font':self.new_font_style(font_name,font_size,False,False,False,default_color),'paragraph':self.new_paragraph_style(2,2)}
        self.add_param(style_name,'name',name_style)
        self.add_param(style_name,'speciality',speciality_style)
        self.add_param(style_name,'contacts',contacts_style)
        self.add_param(style_name,'links',link_style)
        self.add_param(style_name,'section_name',section_name_style)
        self.add_param(style_name,'paragraph_name',paragraph_name_style)
        self.add_param(style_name,'paragraph_text',paragraph_text_style)

    def set_current_style(self,current_style):
        self.current_style=current_style

    def __init__(self,export_file_name='resume.docx'):
        self.document=docx.Document()
        self.styles={}
        self.param=['name','speciality','contacts','section_name','paragraph_name','paragraph_text']
        self.default_style()
        self.current_style=''
        self.set_current_style('default')
        self.name=''
        self.speciality=''
        self.contacts=[]
        self.sections=[]
        self.export_file_name=export_file_name

    def add_name(self,name,speciality):
        self.name=name
        self.speciality=speciality

    #add contact (name,address)
    def add_contact(self,contact_name,contact_address):
        self.contacts.append({'contact_name':contact_name,'contact_address':contact_address})
        
    #return new section
    def new_section(self,name):
        return {'section_name':name,'paragraphs':[]}

    #return new paragraph
    def new_paragraph(self,name,text):
        return {'paragraph_name':name,'paragraph_text':text}
        
    #get section index by name
    def get_section(self,section_name):
        r=-1
        l=len(self.sections)
        if l>0:
            for i in range(l):
                if self.sections[i]['section_name']==section_name:
                    r=i
                    break
        return r
    
    #get paragraph index by name
    def get_paragraph(self,section_index,paragraph_name):
        r=-1
        if section_index!=-1:
            l=len(self.sections[section_index]['paragraphs'])
            if l>0:
                for i in range(l):
                    if self.sections[section_index]['paragraphs'][i]['paragraph_name']==paragraph_name:
                        r=i
                        break
        return r
    
    #add paragraph
    def add_paragraph(self,section_name,paragraph_name,paragraph_text):
        i=self.get_section(section_name)
        if i==-1:
            self.sections.append(self.new_section(section_name))
            i=len(self.sections)-1
        j=self.get_paragraph(i,paragraph_name)
        if j==-1:
            self.sections[i]['paragraphs'].append(self.new_paragraph(paragraph_name,paragraph_text))
        else:
            self.sections[i]['paragraphs'][j]['paragraph_text']=paragraph_text
    
    #apply styles
    def setstyle(self,paragraph,run,name):
        style=self.styles[self.current_style][name]
        font = run.font
        paragraph.paragraph_format.space_before = Pt(style['paragraph']['space_before'])
        paragraph.paragraph_format.space_after = Pt(style['paragraph']['space_after'])
        font_style = style['font']
        font.name = font_style['name']
        font.size = Pt(font_style['size'])
        font.color.rgb = RGBColor(font_style['color'][0],font_style['color'][1],font_style['color'][2])
        run.bold = font_style['bold']
        run.italic = font_style['italic']
        run.underline = font_style['underline']

    def w(self,n,**p):
        xml=docx.oxml.shared.OxmlElement('w:'+n)
        for key,value in p.items():
            xml.set(docx.oxml.shared.qn('w:'+key), value)
        return xml

    def add_hyperlink(self,paragraph,url,text,color,font_size,font_name,font_bold,font_italic,font_underline):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        rPr.append(self.w('color',val=color))
        if font_bold:
            rPr.append(self.w('b',val='true'))
        else:
            rPr.append(self.w('b',val='none'))

        if font_italic:
            rPr.append(self.w('i',val='true'))
        else:
            rPr.append(self.w('i',val='none'))  

            rPr.append(self.w('u',val=font_underline))
        

        f = self.w('rFonts',ascii=font_name,hAnsi=font_name,cs=font_name,eastAsia=font_name)
        rPr.append(f)
        f = self.w('sz',val=str(font_size*2))
        rPr.append(f)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    #render name and speciality
    def render_name(self):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(self.name)
        self.setstyle(paragraph,run,'name')
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(self.speciality)
        self.setstyle(paragraph,run,'speciality')
  
    #render contscts
    def render_contact(self,contact_name,contact_address):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(contact_name+': ')
        self.setstyle(paragraph,run,'contacts')
        # color = '729fcf'
        link_style=self.styles[self.current_style]['links']['font']
        link_color=rgb_to_hex(link_style['color'][0],link_style['color'][1],link_style['color'][2])
        self.add_hyperlink(paragraph,contact_address,contact_address.replace("mailto:",""),link_color,link_style['size'],link_style['name'],False,False,'single')        
    
    def render_contacts(self):
        for contact in self.contacts:
            self.render_contact(contact['contact_name'],contact['contact_address'])

    #render section name
    def render_section_name(self,text):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(text)
        self.setstyle(paragraph,run,'section_name')
        
    #render section paragraph name
    def render_paragraph_name(self,text):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(text)
        self.setstyle(paragraph,run,'paragraph_name')
        
    #render section paragraph text
    def render_paragraph_text(self,text):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(text)
        self.setstyle(paragraph,run,'paragraph_text')

    #render resume
    def render(self):
        self.render_name()
        self.render_contacts()
        self.document.add_paragraph()
        for section in self.sections:
            self.render_section_name(section['section_name'])
            for paragraph in section['paragraphs']:
                self.render_paragraph_name(paragraph['paragraph_name'])
                self.render_paragraph_text(paragraph['paragraph_text'])
                self.document.add_paragraph()

    #save resume to docx file
    def export(self,file_name):
        if file_name:
            self.document.save(file_name)
        else:
            self.document.save(current_dir+self.file_name)

# myresume=resume()
# myresume.add_name('Name','Company')
# myresume.add_contact('email','email@email.com')
# myresume.add_contact('github','https://github.com')
# myresume.add_paragraph('Section 1','title','text')
# myresume.add_paragraph('Section 1','title2','text2')
# myresume.add_paragraph('Summary','Summary information','text3')
# myresume.render()
# myresume.export(current_dir+'export.docx')
