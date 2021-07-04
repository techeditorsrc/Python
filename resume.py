import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.dml.color import ColorFormat
from docx.shared import Pt
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
import os

current_dir=os.getcwd()+'/'
class resume():
    def __init__(self,export_file_name='resume.docx'):
        self.document=docx.Document()
        self.styles={}
        self.name=''
        self.speciality=''
        self.contacts=[]
        self.sections=[]
        self.export_file_name=export_file_name
        
    def add_key(self,obj,key,value):
        if not key in obj:
            obj[key]=value

    def add_name(self,name,speciality):
        self.name=name
        self.speciality=speciality

    #add contact (name,address)
    def add_contact(self,contact_name,contact_address):
        self.contacts.append({'contact_name':contact_name,'contact_address':contact_address})
        
    #return new section
    def new_section(self,name):
        return {'section_name':name,'paragraphs':[]}

    #return new paragraph
    def new_paragraph(self,name,text):
        return {'paragraph_name':name,'paragraph_text':text}
        
    #get section index by name
    def get_section(self,section_name):
        r=-1
        l=len(self.sections)
        if l>0:
            for i in range(l):
                if self.sections[i]['section_name']==section_name:
                    r=i
                    break
        return r
    
    #get paragraph index by name
    def get_paragraph(self,section_index,paragraph_name):
        r=-1
        if section_index!=-1:
            l=len(self.sections[section_index]['paragraphs'])
            if l>0:
                for i in range(l):
                    if self.sections[section_index]['paragraphs'][i]['paragraph_name']==paragraph_name:
                        r=i
                        break
        return r
    
    #add paragraph
    def add_paragraph(self,section_name,paragraph_name,paragraph_text):
        i=self.get_section(section_name)
        if i==-1:
            self.sections.append(self.new_section(section_name))
            i=len(self.sections)-1
        j=self.get_paragraph(i,paragraph_name)
        if j==-1:
            self.sections[i]['paragraphs'].append(self.new_paragraph(paragraph_name,paragraph_text))
        else:
            self.sections[i]['paragraphs'][j]['paragraph_text']=paragraph_text
    
    def setstyle(self,paragraph,run,name):
        font = run.font
        font_size=10
        font.size=Pt(font_size)
        font.name='Dejavu Sans'
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        if name=='name':
            run.bold=True
            font.size=Pt(font_size)
        if name=='speciality':
            font.size=Pt(font_size)
        if name=='contacts':
            font.size=Pt(font_size-4)
        if name=='section_name':
            font.size=Pt(font_size)
            font.color.rgb = RGBColor(65, 138, 183)
        if name=='paragraph_name':
            font.size=Pt(font_size)
            font.color.rgb = RGBColor(53, 108, 132)
        if name=='paragraph_text':
            font.size=Pt(font_size)

    #render name and speciality
    def render_name(self):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(self.name)
        self.setstyle(paragraph,run,'name')
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(self.speciality)
        self.setstyle(paragraph,run,'speciality')
  
    #render contscts
    def render_contact(self,contact_name,contact_address):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(contact_name+': '+contact_address)
        self.setstyle(paragraph,run,'contacts')
        
    def render_contacts(self):
        for contact in self.contacts:
            self.render_contact(contact['contact_name'],contact['contact_address'])

    #render section name
    def render_section_name(self,text):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(text)
        self.setstyle(paragraph,run,'section_name')
        
    #render section paragraph name
    def render_paragraph_name(self,text):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(text)
        self.setstyle(paragraph,run,'paragraph_name')
        
    #render section paragraph text
    def render_paragraph_text(self,text):
        paragraph = self.document.add_paragraph()
        run=paragraph.add_run(text)
        self.setstyle(paragraph,run,'paragraph_text')

    #render resume
    def render(self):
        self.render_name()
        self.render_contacts()
        self.document.add_paragraph()
        for section in self.sections:
            self.render_section_name(section['section_name'])
            for paragraph in section['paragraphs']:
                self.render_paragraph_name(paragraph['paragraph_name'])
                self.render_paragraph_text(paragraph['paragraph_text'])
                self.document.add_paragraph()

    #save resume to docx file
    def export(self,file_name):
        if file_name:
            self.document.save(file_name)
        else:
            self.document.save(current_dir+self.file_name)

myresume=resume()
myresume.add_name('Name','Speciality')
myresume.add_contact('email','email@email.com')
myresume.add_contact('github','https://github.com')
myresume.add_paragraph('Section 1','title','text')
myresume.add_paragraph('Section 1','title2','text2')
myresume.add_paragraph('Summary','Summary information','text3')
myresume.render()
myresume.export(current_dir+'export.docx')
